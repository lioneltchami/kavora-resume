#!/usr/bin/env python3
"""Generate Lucy Hastings' resume as a professionally formatted DOCX."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

NAVY = RGBColor(0x1E, 0x2A, 0x3A)
GOLD = RGBColor(0xB0, 0x8D, 0x57)
DARK = RGBColor(0x2C, 0x2C, 0x2C)
GRAY = RGBColor(0x55, 0x55, 0x55)
FONT_NAME = "Calibri"  # Universal fallback; DM Sans may not be installed

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

style = doc.styles["Normal"]
style.font.name = FONT_NAME
style.font.size = Pt(9.5)
style.font.color.rgb = DARK
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = Pt(12.5)


def add_centered_text(text, size=22, bold=True, color=NAVY, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = FONT_NAME
    return p


def add_gold_line():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    # Create a thin gold line using a border
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="1" w:color="B08D57"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_section_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    run.font.name = FONT_NAME
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="1E2A3A"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def add_paragraph(text, size=9.5, color=DARK, bold=False, italic=False,
                   space_before=0, space_after=0, alignment=None):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = FONT_NAME
    run.bold = bold
    run.italic = italic
    return p


def add_job_header(title, dates):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    # Title
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = NAVY
    run.font.name = FONT_NAME
    # Tab + dates (right-aligned via tab stop)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.3), alignment=2)  # RIGHT
    run2 = p.add_run(f"\t{dates}")
    run2.font.size = Pt(8.5)
    run2.font.color.rgb = GRAY
    run2.font.name = FONT_NAME
    return p


def add_company(text):
    return add_paragraph(text, size=9, color=GRAY, bold=False, space_after=1)


def add_project(text):
    return add_paragraph(text, size=8.5, color=GRAY, italic=True, space_after=2)


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = DARK
    run.font.name = FONT_NAME
    return p


# ═══════════════════════════════════════════════════
# BUILD THE DOCUMENT
# ═══════════════════════════════════════════════════

# HEADER
add_centered_text("LUCY HASTINGS", size=22, bold=True, color=NAVY, space_after=2)
add_centered_text(
    "303-928-9994  |  hotrodlucy@gmail.com  |  Denver, CO & Amarillo, TX",
    size=9.5, bold=False, color=GRAY, space_after=0
)
add_gold_line()

# PROFESSIONAL SUMMARY
add_section_title("Professional Summary")
add_paragraph(
    "Results-driven Construction Superintendent with 25+ years of experience delivering "
    "commercial, residential, government, and industrial projects ranging from $100K to $45M. "
    "Proven track record managing DOE-regulated nuclear facility demolitions, federal installations "
    "including Peterson Air Force Base and the US Mint, and large-scale historical renovations "
    "with zero lost-time incidents. Bilingual leader who combines deep field expertise with "
    "OSHA 30 certification and a 14-year track record as a successful construction business owner.",
    size=9.5, space_after=2
)

# CORE COMPETENCIES
add_section_title("Core Competencies")
competencies = [
    ["Project Scheduling & Coordination", "Safety Program Management (OSHA)", "Quality Assurance & Control"],
    ["Subcontractor & Crew Management", "Budget Forecasting & Cost Control", "DOE & Government Compliance"],
    ["RFI / Submittal / Change Order Mgmt", "Risk Mitigation & Analysis", "Blueprint & Specification Reading"],
    ["Bilingual Communication (EN/ES)", "Contract Negotiation & Procurement", "Construction Document Control"],
]
table = doc.add_table(rows=4, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(competencies):
    for j, cell_text in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(cell_text)
        run.font.size = Pt(8.5)
        run.font.color.rgb = DARK
        run.font.name = FONT_NAME
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        # Remove cell borders
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(tcBorders)

# CERTIFICATIONS
add_section_title("Certifications & Training")
add_paragraph(
    "OSHA 30-Hour  |  OSHA 10-Hour  |  Fall Protection  |  Confined Space Entry  |  "
    "COR  |  Forklift Operation  |  Heavy Equipment Operation  |  Scaffolding Safety  |  "
    "First Aid / CPR / AED  |  Trenching & Shoring  |  LOTO  |  Respirator Fit  |  "
    "Fire Extinguisher  |  GERT  |  Excavation Hazard Mgmt  |  OSHA Accident Investigation",
    size=8.5, color=GRAY, space_after=2
)

# PROFESSIONAL EXPERIENCE
add_section_title("Professional Experience")

# ── Frontera ──
add_job_header("Superintendent & Safety Coordinator", "Aug 2023 – Jul 2025")
add_company("Frontera Construction — Del Rio, TX")
add_project("Mazda Dealership Remodel/Addition — $8M+")
add_bullet("Directed an $8M+ dealership remodel/addition from preconstruction through closeout, coordinating 12+ subcontractor trades and maintaining schedule adherence within 2% of baseline.")
add_bullet("Orchestrated daily field operations including crew deployment, material staging, and equipment logistics while serving as the primary liaison to the architect and owner\u2019s representative.")
add_bullet("Delivered weekly 3-week look-ahead schedules, progress photo documentation, and daily field reports through ProCore, keeping all stakeholders aligned on milestones and change orders.")
add_bullet("Spearheaded on-site safety program including daily toolbox talks, new-hire orientations, and OSHA inspection readiness, achieving zero recordable incidents across the project lifecycle.")
add_bullet("Streamlined RFI and submittal workflows, processing 50+ submittals and 30+ RFIs to maintain uninterrupted construction sequencing.")

# ── Commonwealth ──
add_job_header("Superintendent & Safety Coordinator", "May 2022 – Apr 2023")
add_company("Commonwealth Construction — Fond du Lac, WI")
add_project("St. Anthony\u2019s Hospital Historical Remodel, Senior Living — $45M")
add_bullet("Managed a $45M historical renovation of St. Anthony\u2019s Hospital into a senior citizen living community, overseeing 20+ subcontractor crews across phased interior and exterior scopes.")
add_bullet("Enforced SWPP controls, environmental safety protocols, and harassment prevention policies on a sensitive occupied-building renovation with active community interfaces.")
add_bullet("Coordinated utility tie-ins for electrical, gas, telecom, and cable services across a multi-building campus, eliminating service disruptions during the transition.")
add_bullet("Maintained quality control checkpoints at every construction phase, resulting in zero failed municipal inspections and on-time substantial completion.")

# ── Sames ──
add_job_header("Superintendent & Safety Coordinator", "Feb 2022 – 2022")
add_company("Sames, Inc. — Amarillo, TX")
add_project("Pantex Nuclear Warhead Plant — Demolition — $12M")
add_bullet("Led a $12M demolition project at the Pantex Nuclear Warhead Plant, ensuring full compliance with DOE orders, facility security requirements, and QA certification standards.")
add_bullet("Managed cross-functional project teams including systems engineers, architects, and DOE client staff while maintaining scope, schedule, and budget on a high-security federal site.")
add_bullet("Developed risk mitigation strategies through formal risk identification and response planning, collaborating with upper management to reduce project exposure.")
add_bullet("Embedded lessons-learned protocols across all project staff, driving continuous process improvement on subsequent DOE-regulated scopes.")

# ── B&M ──
add_job_header("Superintendent & Quality Control", "Dec 2021 – Feb 2022")
add_company("B&M Assets — Amarillo, TX")
add_project("New Residential Home Builds — up to $1M")
add_bullet("Supervised construction of new residential homes valued up to $1M, managing 6\u20138 direct reports and multiple subcontractor crews on concurrent builds.")
add_bullet("Conducted technical safety and health reviews identifying corrective actions, maintaining full OSHA compliance across all active job sites.")
add_bullet("Coordinated all building department inspections, permit acquisitions, and utility installations, achieving zero schedule delays from regulatory holds.")

# ── Tri-State ──
add_job_header("Safety Supervisor (SSHO), QA/QC & Backup Superintendent", "Jul 2021 – Oct 2021")
add_company("Tri-State General Contractors — Amarillo, TX")
add_project("16 AISD School Projects ($13M+) & Pantex Demolition ($15M+)")
add_bullet("Oversaw safety compliance on 16 simultaneous school renovation projects for Amarillo ISD totaling $13M+, establishing standardized safety protocols across all campuses.")
add_bullet("Served as SSHO on $15M+ Pantex Nuclear Warhead Plant demolition projects (12R24E, 12R24S, 12R30S), ensuring DOE order compliance and zero safety violations.")
add_bullet("Managed QA/QC programs across all active projects, implementing inspection checklists and corrective action tracking that reduced punch list items by 25%.")

# ── WBI ──
add_job_header("General Superintendent & Safety Coordinator", "Jan 2021 – Jun 2021")
add_company("WBI General Contractors — Amarillo, TX")
add_project("Residential & Commercial Portfolio — $1.2M combined")
add_bullet("Directed a portfolio of residential and commercial projects totaling $1.2M including sunroom additions, full remodels, and roof replacements, delivering all scopes on time and within budget.")
add_bullet("Supervised 6\u20138 field personnel and subcontractors daily, conducting weekly safety meetings and enforcing OSHA requirements across all active job sites.")
add_bullet("Managed end-to-end project execution from estimating and contract negotiation through final inspections and client handoff.")

# ── Clover Leaf ──
add_job_header("PM & VP Assistant / Field Assistant / HR Coordinator", "2019 – 2020")
add_company("Clover Leaf Solutions — Albuquerque, NM")
add_project("Pantex Nuclear Warhead Plant — Zone 4, Mock Mixer, 11-15a Demo — $49M+ combined")
add_bullet("Coordinated project administration for three concurrent Pantex projects valued at $49M+, managing submittals, RFIs, FCRs, and change orders.")
add_bullet("Orchestrated weekly multi-party coordination meetings with Pantex officials, subcontractors, and project personnel across geographically dispersed teams.")
add_bullet("Administered new-hire badging, security clearance coordination, and contractor training programs for DOE-regulated facility access.")

# ── Your Denver Metro ──
add_job_header("General Superintendent / Owner-Operator", "2007 – 2021 (14 years)")
add_company("Your Denver Metro Construction, LLC — Denver, CO")
add_project("Commercial, Residential & Agricultural — Lucent Technology, Capitol Steel, City of Broomfield")
add_bullet("Founded and operated a full-service construction firm for 14 years, delivering $2M+ in projects including Lucent Technology ($500K+), Capitol Steel agricultural builds, and City of Broomfield municipal work ($300K+).")
add_bullet("Managed all business operations including P&L oversight, AR/AP, payroll processing, contract negotiation, and HR administration for a team of 6\u20138 employees plus subcontractors.")
add_bullet("Directed field operations across diverse scopes \u2014 new construction, tenant finishes, fix-and-flips, excavation, concrete, structural steel welding, and metal/wood framing \u2014 maintaining zero lost-time incidents.")
add_bullet("Established long-term client relationships with real estate agents, commercial property managers, and municipal entities, generating consistent repeat business.")

# ── Zap Engineering ──
add_job_header("Assistant Project Manager", "2009 – 2011")
add_company("Zap Engineering — Lakewood, CO")
add_project("Oil & Gas Compressor Station Installations — $37M+ combined")
add_bullet("Supported delivery of two oil & gas compressor station projects valued at $37M+ including Solar Taurus 70 turbine booster compressor installations.")
add_bullet("Maintained master drawing lists across 10+ engineering disciplines and managed all client transmittals and documentation.")
add_bullet("Coordinated multi-disciplinary engineering teams and led weekly status meetings, producing forecasting reports, budget analyses, and resource allocation plans.")

# ── Rey's Welding ──
add_job_header("General Superintendent / Safety Manager / Controller", "1991 – 2006 (15 years)")
add_company("Rey\u2019s Welding & Construction, LLC — Loveland, CO")
add_project("Peterson AFB, US Mint, Arapahoe Justice Center, Fort Lupton Rail — $29M+ combined")
add_bullet("Delivered high-profile federal projects including Peterson Air Force Base addition ($750K+), US Mint vault expansion ($500K+), and Arapahoe Justice Center jail addition ($2.5M+).")
add_bullet("Managed daily operations for 8\u201310 employees and subcontractors on projects totaling $29M+ including Fort Lupton Rail Vacation Cars ($20M+) and commercial contracts ($5M+).")
add_bullet("Oversaw full financial operations including budgeting, AR/AP, payroll, workers\u2019 compensation, and P&L reporting for a multi-million-dollar construction firm.")

# EDUCATION
add_section_title("Education")
p = add_paragraph("", space_after=1)
run = p.add_run("A.A.S. Degree")
run.bold = True
run.font.size = Pt(9)
run.font.color.rgb = NAVY
run.font.name = FONT_NAME
run2 = p.add_run("  \u2014  Aims Community College, Greeley, CO  |  GPA: 3.8\u20134.0")
run2.font.size = Pt(9)
run2.font.color.rgb = DARK
run2.font.name = FONT_NAME

p = add_paragraph("", space_after=1)
run = p.add_run("Coursework")
run.bold = True
run.font.size = Pt(9)
run.font.color.rgb = NAVY
run.font.name = FONT_NAME
run2 = p.add_run("  \u2014  Amarillo College, Amarillo, TX  |  GPA: 3.8")
run2.font.size = Pt(9)
run2.font.color.rgb = DARK
run2.font.name = FONT_NAME

p = add_paragraph("", space_after=1)
run = p.add_run("High School Diploma")
run.bold = True
run.font.size = Pt(9)
run.font.color.rgb = NAVY
run.font.name = FONT_NAME
run2 = p.add_run("  \u2014  Hereford High School, Hereford, TX  |  GPA: 3.8  |  Dean\u2019s Honor List")
run2.font.size = Pt(9)
run2.font.color.rgb = DARK
run2.font.name = FONT_NAME

# Save
output_path = os.path.join(os.path.dirname(__file__), "Lucy_Hastings_Resume.docx")
doc.save(output_path)
print(f"DOCX saved to: {output_path}")
